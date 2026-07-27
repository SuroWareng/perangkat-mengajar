from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
import json
import os
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

app = Flask(__name__)
CORS(app)

# Folder untuk menyimpan data
DATA_FOLDER = 'data'
if not os.path.exists(DATA_FOLDER):
    os.makedirs(DATA_FOLDER)

# ==================== ROUTES ====================

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/prota', methods=['GET', 'POST'])
def prota():
    if request.method == 'POST':
        data = request.json
        filename = f"{DATA_FOLDER}/prota_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return jsonify({"status": "success", "message": "Data Prota disimpan", "file": filename})
    
    return jsonify({"status": "error"})

@app.route('/api/promes', methods=['GET', 'POST'])
def promes():
    if request.method == 'POST':
        data = request.json
        filename = f"{DATA_FOLDER}/promes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return jsonify({"status": "success", "message": "Data Promes disimpan", "file": filename})
    
    return jsonify({"status": "error"})

@app.route('/api/modul', methods=['GET', 'POST'])
def modul():
    if request.method == 'POST':
        data = request.json
        filename = f"{DATA_FOLDER}/modul_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return jsonify({"status": "success", "message": "Data Modul disimpan", "file": filename})
    
    return jsonify({"status": "error"})

# ==================== Reference & Template Endpoints ====================

@app.route('/api/reference/cp-fase', methods=['GET'])
def reference_cp_fase():
    """Return CP-per-phase reference JSON (docs/cp_fase_2026.json).
    This provides a structured summary of Fase A–E untuk Kurikulum 2026.
    """
    try:
        with open('docs/cp_fase_2026.json', 'r', encoding='utf-8') as f:
            data = json.load(f)
        return jsonify({"status": "success", "data": data})
    except FileNotFoundError:
        return jsonify({"status": "error", "message": "Reference file not found. Please ensure docs/cp_fase_2026.json exists."}), 404

@app.route('/api/templates/prota', methods=['GET'])
def template_prota():
    """Return a JSON template for Prota (Program Tahunan) to be used by the front-end or for export tests."""
    template = {
        "identitas": {
            "nama_sekolah": "",
            "alamat": "",
            "kota": "",
            "provinsi": "",
            "mata_pelajaran": "",
            "kelas": "",
            "tahun_ajaran": "",
            "nama_guru": "",
            "nip": ""
        },
        "profil_pelajar_pancasila": [],
        "pembelajaran": [
            {"elemen": "", "cp": "", "jam": 0, "p5": ""}
        ],
        "proyek_p5": {"tema": "", "alokasi_jam": 0, "dimensi_p5": "", "output": ""},
        "keterangan": ""
    }
    return jsonify({"status": "success", "template": template})

@app.route('/api/templates/promes', methods=['GET'])
def template_promes():
    template = {
        "identitas": {
            "nama_sekolah": "",
            "mata_pelajaran": "",
            "kelas": "",
            "semester": "",
            "tahun_ajaran": ""
        },
        "minggu": [
            {"minggu_ke": 1, "tanggal": "", "tujuan": "", "materi": "", "kegiatan": "", "jp": 0, "penilaian": "", "sumber": "", "keterangan": ""}
        ]
    }
    return jsonify({"status": "success", "template": template})

@app.route('/api/templates/modul', methods=['GET'])
def template_modul():
    template = {
        "identitas": {"judul_modul": "", "nama_penyusun": "", "sekolah": "", "kelas": "", "mata_pelajaran": "", "alokasi_waktu": ""},
        "capaian_pembelajaran": "",
        "tujuan_pembelajaran": "",
        "pemahaman_konsep": "",
        "pemahaman_relevansi": "",
        "pertanyaan_pemantik": "",
        "kegiatan_pembelajaran": "",
        "penguatan_p5": "",
        "penilaian_formatif": "",
        "penilaian_sumatif": "",
        "refleksi": "",
        "lkpd": "",
        "sumber_belajar": ""
    }
    return jsonify({"status": "success", "template": template})

# ==================== EXPORT FUNCTIONS ====================

@app.route('/api/export/prota-word-2026', methods=['POST'])
def export_prota_word_2026():
    """Export Prota 2026 ke format Word"""
    data = request.json
    
    doc = Document()
    
    # Header
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('PROGRAM TAHUNAN (PROTA) - KURIKULUM 2026')
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Identitas Sekolah
    doc.add_heading('I. IDENTITAS SEKOLAH', level=2)
    table = doc.add_table(rows=9, cols=2)
    table.style = 'Light Grid Accent 1'
    
    identitas = [
        ('Nama Sekolah', data.get('nama_sekolah', '')),
        ('Alamat', data.get('alamat', '')),
        ('Kota', data.get('kota', '')),
        ('Provinsi', data.get('provinsi', '')),
        ('Mata Pelajaran', data.get('mata_pelajaran', '')),
        ('Kelas', data.get('kelas', '')),
        ('Tahun Ajaran', data.get('tahun_ajaran', '')),
        ('Nama Guru', data.get('nama_guru', '')),
        ('NIP', data.get('nip', ''))
    ]
    
    for i, (label, value) in enumerate(identitas):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    
    # Profil Pelajar Pancasila (P5)
    doc.add_heading('II. PROFIL PELAJAR PANCASILA', level=2)
    p5_list = data.get('profil_pelajar_pancasila', [])
    for p5 in p5_list:
        doc.add_paragraph(f"{p5}", style='List Bullet')
    
    # Elemen & Capaian Pembelajaran
    doc.add_heading('III. ELEMEN PEMBELAJARAN & CAPAIAN PEMBELAJARAN (CP)', level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    headers = ['No', 'Elemen', 'Capaian Pembelajaran (CP)', 'Alokasi Waktu (JP)', 'P5 Terkait']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data pembelajaran
    pembelajaran = data.get('pembelajaran', [])
    for i, item in enumerate(pembelajaran, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = item.get('elemen', '')
        row_cells[2].text = item.get('cp', '')
        row_cells[3].text = str(item.get('jam', ''))
        row_cells[4].text = item.get('p5', '')
    
    # Proyek P5
    doc.add_heading('IV. PROYEK PENGUATAN PROFIL PELAJAR PANCASILA', level=2)
    proyek_p5 = data.get('proyek_p5', {})
    doc.add_paragraph(f"Tema: {proyek_p5.get('tema', '')}")
    doc.add_paragraph(f"Alokasi Waktu: {proyek_p5.get('alokasi_jam', '')} JP")
    doc.add_paragraph(f"Dimensi P5: {proyek_p5.get('dimensi_p5', '')}")
    doc.add_paragraph(f"Output/Deliverable: {proyek_p5.get('output', '')}")
    
    # Total alokasi waktu
    doc.add_heading('V. TOTAL ALOKASI WAKTU', level=2)
    total_jam = sum([int(item.get('jam', 0)) for item in pembelajaran])
    jam_proyek = int(proyek_p5.get('alokasi_jam', 0))
    doc.add_paragraph(f"Intrakurikuler: {total_jam} Jam Pelajaran")
    doc.add_paragraph(f"Kokurikuler (Proyek P5): {jam_proyek} Jam Pelajaran")
    doc.add_paragraph(f"Total: {total_jam + jam_proyek} Jam Pelajaran")
    
    # Keterangan
    doc.add_paragraph(f"Keterangan: {data.get('keterangan', '')}")
    
    # Tanda tangan
    doc.add_page_break()
    doc.add_paragraph()
    doc.add_paragraph("Mengetahui,")
    table = doc.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "Kepala Sekolah"
    table.rows[0].cells[1].text = "Guru Mata Pelajaran"
    table.rows[3].cells[0].text = f"NIP. {data.get('nip_kepsek', '')}"
    table.rows[3].cells[1].text = f"NIP. {data.get('nip', '')}"
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"Prota_2026_{data.get('mata_pelajaran', 'Pelajaran')}_{datetime.now().strftime('%Y%m%d')}.docx"
    )

@app.route('/api/export/modul-word-2026', methods=['POST'])
def export_modul_word_2026():
    """Export Modul Ajar 2026 ke format Word"""
    data = request.json
    
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('MODUL AJAR - KURIKULUM 2026')
    run.font.size = Pt(14)
    run.font.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(data.get('judul_modul', ''))
    run.font.size = Pt(12)
    run.font.bold = True
    
    # A. INFORMASI UMUM
    doc.add_heading('A. INFORMASI UMUM', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    informasi = [
        ('Nama Penyusun', data.get('nama_penyusun', '')),
        ('Sekolah', data.get('sekolah', '')),
        ('Kelas', data.get('kelas', '')),
        ('Elemen Pembelajaran', data.get('elemen', '')),
        ('Capaian Pembelajaran (CP)', data.get('cp', '')),
        ('Dimensi P5', data.get('dimensi_p5', '')),
        ('Alokasi Waktu', f"{data.get('alokasi_waktu', '')} Jam Pelajaran"),
        ('Mata Pelajaran', data.get('mata_pelajaran', '')),
    ]
    
    for label, value in informasi:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
    
    # B. KOMPONEN INTI
    doc.add_heading('B. KOMPONEN INTI', level=2)
    
    doc.add_heading('1. Tujuan Pembelajaran', level=3)
    doc.add_paragraph(data.get('tujuan_pembelajaran', ''), style='List Bullet')
    
    doc.add_heading('2. Pemahaman Bermakna', level=3)
    doc.add_paragraph(f"Konsep Utama: {data.get('pemahaman_konsep', '')}")
    doc.add_paragraph(f"Relevansi: {data.get('pemahaman_relevansi', '')}")
    
    doc.add_heading('3. Pertanyaan Pemantik', level=3)
    pertanyaan_list = data.get('pertanyaan_pemantik', '').split('\n')
    for q in pertanyaan_list:
        if q.strip():
            doc.add_paragraph(q.strip(), style='List Bullet')
    
    doc.add_heading('4. Kegiatan Pembelajaran', level=3)
    doc.add_paragraph(data.get('kegiatan_pembelajaran', ''))
    
    doc.add_heading('5. Penguatan Profil Pelajar Pancasila', level=3)
    doc.add_paragraph(data.get('penguatan_p5', ''))
    
    doc.add_heading('6. Asesmen/Penilaian', level=3)
    doc.add_paragraph(f"Formatif: {data.get('penilaian_formatif', '')}")
    doc.add_paragraph(f"Sumatif: {data.get('penilaian_sumatif', '')}")
    
    doc.add_heading('7. Refleksi Peserta Didik', level=3)
    doc.add_paragraph(data.get('refleksi', ''))
    
    # C. LAMPIRAN
    doc.add_heading('C. LAMPIRAN', level=2)
    doc.add_heading('1. Lembar Kerja Peserta Didik (LKPD)', level=3)
    doc.add_paragraph(data.get('lkpd', ''))
    
    doc.add_heading('2. Sumber Belajar', level=3)
    sumber_list = data.get('sumber_belajar', '').split('\n')
    for s in sumber_list:
        if s.strip():
            doc.add_paragraph(s.strip(), style='List Bullet')
    
    # Save
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"Modul_2026_{data.get('judul_modul', 'Ajar')}_{datetime.now().strftime('%Y%m%d')}.docx"
    )

# Keep old endpoints for compatibility
@app.route('/api/export/prota-word', methods=['POST'])
def export_prota_word():
    return export_prota_word_2026()

@app.route('/api/export/modul-word', methods=['POST'])
def export_modul_word():
    return export_modul_word_2026()

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
